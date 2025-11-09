#!/usr/bin/env python3
"""
Crea un DOCX minimale con due paragrafi per verificare lo spacing verticale.
"""

from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_paragraph("Line 1 - Lorem ipsum").runs[0].font.size = Pt(11)
doc.add_paragraph("Line 2 - Lorem ipsum").runs[0].font.size = Pt(11)

doc.save("samples/simple-spacing.docx")
print("✓ File creato: samples/simple-spacing.docx")
