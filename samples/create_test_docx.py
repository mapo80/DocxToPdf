#!/usr/bin/env python3
"""
Crea un file DOCX di test con Lorem Ipsum per testare il converter DOCX → PDF.
"""

from docx import Document
from docx.shared import Inches, Pt

# Crea documento
doc = Document()

# Imposta margini (1 inch = 72pt su tutti i lati - default Word)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width

# Titolo
title = doc.add_paragraph('Lorem Ipsum Test Document')
title_run = title.runs[0]
title_run.font.size = Pt(24)
title_run.bold = True

# Paragrafo 1
p1 = doc.add_paragraph(
    'Lorem ipsum dolor sit amet, consectetur adipiscing elit. '
    'Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. '
    'Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris '
    'nisi ut aliquip ex ea commodo consequat.'
)
p1.runs[0].font.size = Pt(11)

# Paragrafo 2
p2 = doc.add_paragraph(
    'Duis aute irure dolor in reprehenderit in voluptate velit esse cillum '
    'dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non '
    'proident, sunt in culpa qui officia deserunt mollit anim id est laborum.'
)
p2.runs[0].font.size = Pt(11)

# Paragrafo 3 con mix di stili
p3 = doc.add_paragraph()
r1 = p3.add_run('Testo normale. ')
r1.font.size = Pt(11)

r2 = p3.add_run('Testo grassetto. ')
r2.font.size = Pt(11)
r2.bold = True

r3 = p3.add_run('Testo corsivo. ')
r3.font.size = Pt(11)
r3.italic = True

r4 = p3.add_run('Testo con àccénti e caratteri speciali: €, ©, ™, 👋')
r4.font.size = Pt(11)

# Paragrafo 4 - lungo per testare wrapping multi-riga
p4 = doc.add_paragraph(
    'Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium '
    'doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore '
    'veritatis et quasi architecto beatae vitae dicta sunt explicabo. Nemo enim '
    'ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia '
    'consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.'
)
p4.runs[0].font.size = Pt(11)

# Paragrafo 5
p5 = doc.add_paragraph(
    'At vero eos et accusamus et iusto odio dignissimos ducimus qui blanditiis '
    'praesentium voluptatum deleniti atque corrupti quos dolores et quas molestias '
    'excepturi sint occaecati cupiditate non provident, similique sunt in culpa qui '
    'officia deserunt mollitia animi, id est laborum et dolorum fuga.'
)
p5.runs[0].font.size = Pt(11)

# Salva
doc.save('samples/lorem.docx')
print('✓ File creato: samples/lorem.docx')
