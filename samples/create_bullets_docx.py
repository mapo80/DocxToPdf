#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt


doc = Document()

doc.add_paragraph('Default bullet', style='List Bullet')
doc.add_paragraph('Indented bullet', style='List Bullet 2')
doc.add_paragraph('Another level', style='List Bullet 3')

doc.add_paragraph('Wingdings bullet test').style = 'Normal'
run = doc.paragraphs[-1].runs[0]
run.text = ''
p = doc.add_paragraph(style='List Bullet')
p.add_run('Wingdings style bullet')
p.runs[0].font.name = 'Wingdings'

p = doc.add_paragraph(style='List Number')
p.add_run('Numbered item')

output = 'samples/bullets-basic.docx'
doc.save(output)
print(f'✓ File creato: {output}')
